from __future__ import annotations

import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\AI-Agent学习计划\AI-Agent study")
REFERENCE = Path(r"C:\Users\xys\.codex\plugins\cache\openai-curated-remote\openai-templates\0.1.1\skills\artifact-template-design-report\assets\reference.docx")
OUTPUT = ROOT / "tasks" / "调研模式前端框架重构计划书.docx"
SCREENSHOT = Path(r"C:\Users\xys\AppData\Local\Temp\codex-clipboard-4504ac70-c26b-4670-ad5f-f4eecf5f3bba.png")


def east_asia(run, name="Microsoft YaHei"):
    run.font.name = name
    if run._element.get_or_add_rPr().rFonts is None:
        run._element.get_or_add_rPr().append(OxmlElement("w:rFonts"))
    fonts = run._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:eastAsia"), name)
    fonts.set(qn("w:ascii"), "Helvetica Neue")
    fonts.set(qn("w:hAnsi"), "Helvetica Neue")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9D9D9")
        borders.append(element)


def set_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = text
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "111111")
        if widths:
            set_cell_width(cell, widths[index])
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)
                east_asia(run)
    for r_index, values in enumerate(rows):
        cells = table.add_row().cells
        for c_index, value in enumerate(values):
            cells[c_index].text = str(value)
            cells[c_index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                set_cell_width(cells[c_index], widths[c_index])
            if r_index % 2 == 1:
                set_cell_shading(cells[c_index], "F4F4F4")
            for paragraph in cells[c_index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.08
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
                    east_asia(run)
    doc.add_paragraph("")
    return table


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    set_numbering(p, 1)
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        first.bold = True
        east_asia(first)
        rest = p.add_run(text[len(bold_prefix):])
        east_asia(rest)
    else:
        run = p.add_run(text)
        east_asia(run)
    return p


def add_number(doc, text):
    p = doc.add_paragraph()
    set_numbering(p, 2)
    run = p.add_run(text)
    east_asia(run)
    return p


def add_para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        first.bold = True
        east_asia(first)
        rest = p.add_run(text[len(bold_prefix):])
        east_asia(rest)
    else:
        run = p.add_run(text)
        east_asia(run)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        east_asia(run)
    return p


def remove_content_after_cover(doc):
    body = doc._element.body
    cover_table = doc.tables[0]._element
    reached_cover_table = False
    for child in list(body):
        if child is cover_table:
            reached_cover_table = True
            continue
        if reached_cover_table and child.tag != qn("w:sectPr"):
            # Preserve the paragraph immediately following the cover table: it owns the section break.
            if child.tag == qn("w:p") and child.find(".//w:sectPr", namespaces=child.nsmap) is not None:
                continue
            body.remove(child)


def replace_cover_image(doc):
    if not SCREENSHOT.exists():
        return
    paragraph = doc.paragraphs[0]
    width = Inches(6.5)
    if doc.inline_shapes:
        width = doc.inline_shapes[0].width
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(SCREENSHOT), width=width)


def fill_cover(doc):
    replace_cover_image(doc)
    title = doc.paragraphs[2]
    title.text = "调研模式前端框架重构计划书"
    for run in title.runs:
        east_asia(run)
    cover = doc.tables[0]
    cover.cell(0, 0).text = "保持现有调研链路不变，构建独立双栏研究工作区与双报告阅读体验"
    cover.cell(0, 2).text = "基于现有代码与参考 UI 分析\n2026 年 8 月"
    for cell in cover.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                east_asia(run)


def add_content(doc):
    add_heading(doc, "执行摘要", 1)
    add_para(doc, "本次重构应被定义为一次前端壳层与报告呈现层升级，而不是调研能力重写。现有 /deep_research 请求、三类调研引擎、SSE 事件顺序、NodeProgressPanel 链路翻译、千问反问确认、消息级来源绑定与会话持久化全部保持不变。前端新增独立 ResearchWorkspace：左侧完整承载当前对话流和调研链路，右侧承载“深度研报”和“文本报告”两个互不丢失状态的报告视图。")
    add_para(doc, "AI 写作模式提供可复用的设计经验——独立全屏工作区、可拖拽双栏、右侧文档标题栏、复制/下载/全屏动作、Tiptap 编辑器与 Word 导出——但调研模式不得直接挂载 WritingWorkspace 或共享其业务状态。建议只抽取无业务含义的底层原语，或在 research feature 内实现等价组件，以避免写作与调研互相污染。")
    add_heading(doc, "一览", 2)
    add_table(doc, ["主题", "现状判断", "目标决策"], [
        ["调研链路", "已覆盖 Firecrawl、自研、千问原生及 Agent Loop 事件", "冻结 API 与事件契约，只接入新的视图适配层"],
        ["页面框架", "调研仍位于通用 ChatInterface 单栏消息流", "新增独立 ResearchWorkspace 双栏工作区"],
        ["左侧区域", "消息、节点、来源、反问卡片分散在通用聊天渲染", "完整迁移现有调研对话视图，不省略任何链路"],
        ["右侧区域", "最终 report 仅作为 Markdown 气泡展示", "双标签：深度研报 / 文本报告"],
        ["文档能力", "已有 Tiptap、Word 导出、复制/下载/全屏经验", "research 专属编辑器、导出器和动作栏"],
    ], [1700, 3400, 4260])

    add_heading(doc, "关键发现", 1)
    add_heading(doc, "1. 当前调研模式框架", 2)
    add_bullet(doc, "入口与状态集中在 ChatInterface.tsx：mode='research'、researchEngine、researchOptions、researchChunks、researchProgress、消息与会话恢复均由同一大型组件管理。", "入口与状态集中在 ChatInterface.tsx：")
    add_bullet(doc, "请求契约位于 lib/api.ts：sendDeepResearch 固定调用 POST /deep_research，并把 research_process 翻译为 NodeEvent；最终通过 onResearchReasonDone 与 onResearchDone 接收 report、reasoning、top_chunks 等数据。", "请求契约位于 lib/api.ts：")
    add_bullet(doc, "链路可视化已统一到 NodeProgressPanel：每条 assistant 消息保留独立 nodeProgress、webDocs、researchChunks，支持多轮历史恢复；这是必须原样保留的核心交互资产。", "链路可视化已统一到 NodeProgressPanel：")
    add_bullet(doc, "最终报告仍通过 MarkdownMessage 渲染在消息气泡中，宽度、层级、文档操作和长文阅读体验不足，无法承载参考图所示的专业研报。", "最终报告仍通过 MarkdownMessage 渲染在消息气泡中，")

    add_heading(doc, "2. AI 写作模式可借鉴的框架", 2)
    add_bullet(doc, "WritingWorkspace 已采用独立 fixed 工作区，不受通用聊天内容宽度限制。")
    add_bullet(doc, "工作区使用“左对话 / 分隔条 / 右文档”的网格，左栏比例可在 30%—65% 之间调整。")
    add_bullet(doc, "右侧标题栏已形成标签切换、复制、下载、全屏和文件清单的成熟交互。")
    add_bullet(doc, "WritingDocumentEditor 已接入 Tiptap StarterKit 与 Underline；WritingLayoutWorkspace 具备固定纸张、缩放、分页、模板与 PDF/Word 输出经验。")
    add_para(doc, "关键约束。可借鉴的是页面骨架、交互规则和无业务依赖的视觉原语，不可直接共用 WritingWorkspace、writingDoc、写作时间线或写作导出业务。调研必须有自己的 feature 边界、类型、reducer、持久化字段和报告导出器。", "关键约束。")

    add_heading(doc, "3. 参考 UI 的视觉与交互语言", 2)
    add_bullet(doc, "大面积白色画布、极弱边界、轻阴影与 8—16 px 圆角，整体克制而非卡片堆叠。")
    add_bullet(doc, "左右约 1:1 起始比例，分栏边界清晰；两侧各自滚动，右侧标题栏常驻。")
    add_bullet(doc, "右侧顶部居中双标签，右上角以图标为主的复制、下载、全屏/放大按钮。")
    add_bullet(doc, "深度研报具有封面式标题、元数据、核心摘要深色卡片、章节层级、总结表格和数据图表；文本报告保持规整 Markdown 排版但不额外生成图表。")

    add_heading(doc, "影响与目标框架", 1)
    add_heading(doc, "目标信息架构", 2)
    add_table(doc, ["区域", "职责", "明确不做"], [
        ["ResearchWorkspace 壳", "独立全屏层、双栏布局、分隔条、焦点和滚动管理", "不挂载 WritingWorkspace，不读取 writingDoc"],
        ["左侧 ResearchConversationPane", "完整对话、NodeProgressPanel、来源、反问确认、输入器和引擎选项", "不删减现有事件、不合并多轮状态"],
        ["右侧 ResearchReportPane", "标题、双标签、操作栏、加载/空/错误态、文档滚动", "不参与调研请求调度"],
        ["深度研报", "结构化文档、摘要表、证据表、可验证图表、Tiptap/只读排版、Word/PDF", "无数据时不编造图表"],
        ["文本报告", "原始模型 report 的规范 Markdown/GFM 阅读与复制下载", "不注入深度研报的总结或图表"],
    ], [2100, 4300, 2960])

    add_heading(doc, "建议组件边界", 2)
    add_para(doc, "建议在 frontend/ai-agent/src/features/deep-research/ 下新建完整功能域：")
    for item in [
        "ResearchWorkspace.tsx — 独立工作区入口与三栏网格（左栏、分隔条、右栏）。",
        "researchWorkspaceTypes.ts / researchWorkspaceReducer.ts — 只管理 UI 与报告派生状态。",
        "conversation/ResearchConversationPane.tsx — 从当前聊天渲染中抽出调研专属视图。",
        "report/ResearchReportPane.tsx — 标题栏、标签、加载和错误边界。",
        "report/DeepResearchDocument.tsx — Tiptap 文档或结构化只读渲染器。",
        "report/RawResearchReport.tsx — 原始 report 的规范 Markdown 渲染器。",
        "report/researchReportAdapter.ts — 将 report、top_chunks、pages 转换为安全的展示模型。",
        "report/ResearchChart.tsx / ResearchSummaryTable.tsx — 仅消费可追溯结构化数据。",
        "export/researchWordExport.ts / researchDownload.ts — research 专属 Word、PDF、Markdown、TXT 导出。",
        "components/ResearchReportActions.tsx — 复制、下载、全屏及反馈状态。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "数据流：保持调研链路不变", 2)
    add_number(doc, "ChatInterface 仍负责会话选择、mode 切换和既有 sendDeepResearch 调用。")
    add_number(doc, "所有 onResearchProcess/onNode/onQwenFeedback/onResearchReasonDone/onResearchDone 回调与当前顺序保持不变。")
    add_number(doc, "ResearchWorkspace 通过明确 props 或 research view-model 接收 messages、nodeProgress、researchChunks、report 与操作回调。")
    add_number(doc, "researchReportAdapter 在前端派生 deepReportDocument；rawReport 始终保存原始 report 字符串。")
    add_number(doc, "会话快照新增可选 researchWorkspace 字段，仅保存 UI 状态和派生版本；旧会话缺失该字段时可从最后一条调研 assistant 消息即时重建。")

    add_heading(doc, "深度研报呈现策略", 2)
    add_bullet(doc, "文档模型：标题、日期/主题、核心摘要、章节、证据引用、总结表、图表、来源附录。")
    add_bullet(doc, "Tiptap：复用依赖，不复用写作实例。创建 research 专属 extensions 配置；默认以阅读为主，可按产品决定开放有限编辑。")
    add_bullet(doc, "表格：优先从 Markdown 表格、显式键值对和 top_chunks 元数据生成；每个单元格可追溯至原报告或来源。")
    add_bullet(doc, "图表：只在报告中存在明确、同口径的数字序列时生成；否则显示“暂无可验证图表数据”，绝不以视觉需要补造数据。")
    add_bullet(doc, "导出：深度研报支持格式化 DOCX/PDF；文本报告支持 Markdown/TXT，亦可选原始 DOCX。下载菜单必须明确当前标签与格式。")

    add_heading(doc, "推荐方案", 1)
    add_heading(doc, "实施阶段", 2)
    add_table(doc, ["阶段", "主要工作", "完成标志"], [
        ["P0 契约冻结", "建立现状快照测试；记录事件、消息、会话与引擎契约", "现有调研测试全部通过，新增链路回归基线"],
        ["P1 独立壳层", "新增 ResearchWorkspace、双栏、分隔条、右侧空态与标题栏", "research 模式进入独立工作区，其他模式无变化"],
        ["P2 左栏迁移", "迁移调研消息、节点、来源、反问确认、输入器和选项", "三种引擎全链路与历史多轮均无缺失"],
        ["P3 双报告", "接入 raw report 与 deep report adapter、双标签及稳定滚动", "切换标签不丢内容、位置和选择状态"],
        ["P4 文档增强", "Tiptap、摘要表、证据表、可验证图表、来源脚注", "长文、表格和图表布局符合参考 UI"],
        ["P5 操作与导出", "复制、下载、全屏、Word/PDF/MD/TXT 与失败反馈", "按钮可键盘访问，文件内容与当前标签一致"],
        ["P6 验收上线", "响应式、性能、可访问性、视觉回归、旧会话恢复、灰度开关", "回归门禁通过，可一键回退旧 UI"],
    ], [1050, 4900, 3410])

    add_heading(doc, "关键实施规则", 2)
    add_bullet(doc, "先写契约测试再迁移 UI，尤其锁定 sendDeepResearch 请求体和事件回调顺序。")
    add_bullet(doc, "通过 feature flag 切换旧/新 research UI；新壳只消费既有状态，便于快速回滚。")
    add_bullet(doc, "避免把 ChatInterface 的全部状态复制到新 feature；用 ResearchWorkspaceProps 和 view-model 建立最窄接口。")
    add_bullet(doc, "右栏报告生成采用 memoized adapter 或 Web Worker，避免长报告解析阻塞 SSE 与左侧滚动。")
    add_bullet(doc, "所有图表、表格和引用都保留 sourceIds，复制与导出时同步写入来源说明。")

    add_heading(doc, "验收标准", 2)
    for item in [
        "链路完整：Firecrawl、自研、千问原生、Agent Loop、反问确认和失败降级均与改造前一致。",
        "左侧完整：历史消息、每轮节点、来源、精选片段、reasoning 展开、输入器与选项均未省略。",
        "右侧完整：深度研报与文本报告均可独立滚动、切换、恢复；长文不溢出。",
        "操作完整：复制、下载、全屏/退出全屏均有成功或失败反馈，键盘与屏幕阅读器可用。",
        "文档可信：图表和总结表格均能追溯到 report 或 sources；无可验证数据时不生成伪图。",
        "兼容恢复：旧 research 会话不迁移数据库也能打开；新会话刷新后保持双报告和 UI 状态。",
        "视觉一致：布局、层级、留白、标签和工具栏与参考图同一视觉语言，并与 AI 写作工作区保持产品一致性。",
        "隔离可靠：writing feature 与 deep-research feature 互不导入业务 reducer、document state 或导出器。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "风险与处置", 2)
    add_table(doc, ["风险", "影响", "处置"], [
        ["ChatInterface 过大", "迁移时容易引入状态闭包和竞态", "先提取 view-model/回调，不改请求函数；按阶段切换"],
        ["报告只有字符串", "难以稳定生成图表", "建立可追溯解析规则；不满足规则则降级为摘要表/正文"],
        ["Tiptap 与流式更新冲突", "光标跳动、文档重置", "生成完成后一次装载，或用事务增量更新并保护 selection"],
        ["双栏窄屏", "移动端难同时阅读", "桌面双栏；小屏切换为“对话/报告”顶层页签"],
        ["导出与网页不一致", "用户不信任产物", "统一 ResearchDocument AST 驱动网页与 DOCX/PDF 输出"],
        ["旧会话缺新状态", "历史报告空白", "从 messages 最后一条 research assistant 内容惰性重建"],
    ], [2200, 2900, 4260])

    add_heading(doc, "结论", 2)
    add_para(doc, "推荐采用“冻结调研链路 + 新建独立 ResearchWorkspace + 双报告展示适配层”的路线。它能最大限度复用当前已经稳定的调研能力和 AI 写作积累的交互经验，同时满足业务隔离要求。第一轮应先交付独立双栏、完整左侧链路、右侧文本报告与按钮；第二轮再接入深度研报 AST、Tiptap、可信表格/图表和格式化导出，以降低一次性重构风险。")

    add_heading(doc, "附录", 1)
    add_heading(doc, "A. 现有关键文件", 2)
    add_table(doc, ["文件", "当前职责", "计划动作"], [
        ["components/ChatInterface.tsx", "调研状态、请求回调、通用消息 UI、会话恢复", "保留编排；逐步把 research 视图移至专属 workspace"],
        ["lib/api.ts", "类型、POST /deep_research、SSE 解析与 NodeEvent 翻译", "冻结公开契约，仅补充 UI 所需可选类型"],
        ["components/NodeProgressPanel.tsx", "节点、迭代、来源和 reasoning 展开", "原样嵌入左栏，必要时只做视觉适配"],
        ["components/MarkdownMessage.tsx", "GFM/数学/表格渲染", "用于 raw report 或封装 research 专属样式"],
        ["features/ai-writing/WritingWorkspace.tsx", "独立双栏与文档操作参考", "仅参考，不作为 research 容器"],
        ["features/ai-writing/layout/WritingDocumentEditor.tsx", "Tiptap 接入参考", "复制技术模式，建立 research 专属配置"],
        ["features/ai-writing/thesis/thesisWordExport.ts", "Word 导出参考", "建立 researchWordExport，不共享写作业务模型"],
    ], [3150, 3150, 3060])

    add_heading(doc, "B. 事件兼容矩阵", 2)
    add_table(doc, ["事件/字段", "现有用途", "新 UI 落点"], [
        ["research_process", "阶段进度与节点翻译", "左栏 NodeProgressPanel"],
        ["qwen_feedback", "反问确认与继续研究", "左栏对话内嵌卡片"],
        ["research_reason_done", "reasoning、report、耗时", "左栏 reasoning + 右栏原始报告"],
        ["research done", "total_pages、chunks、top_chunks、report", "左栏来源统计 + 右栏深度研报适配"],
        ["messages[].nodeProgress", "每轮链路持久化", "左栏历史轮次恢复"],
        ["messages[].researchChunks", "每轮精选片段持久化", "来源面板、证据表和报告附录"],
    ], [2200, 3100, 4060])

    add_heading(doc, "C. 本计划明确不包含", 2)
    add_bullet(doc, "不调整后端调研算法、模型、检索、Rerank、推理、降级或 SSE 协议。")
    add_bullet(doc, "不把 AI 写作的 WritingWorkspace 直接复用于调研模式。")
    add_bullet(doc, "不为了生成图表而修改模型输出或伪造数值。")
    add_bullet(doc, "不在本计划阶段实施代码；实际开发应按 P0—P6 分阶段评审。")


def normalize_fonts(doc):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            east_asia(run)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        east_asia(run)


def replace_header_placeholders(path):
    temp = path.with_suffix(".header-patched.docx")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(temp, "w", zipfile.ZIP_DEFLATED) as target:
        for item in source.infolist():
            data = source.read(item.filename)
            if item.filename.startswith("word/header") and item.filename.endswith(".xml"):
                text = data.decode("utf-8")
                text = text.replace("Report title", "调研模式前端框架重构")
                text = text.replace(">Date<", ">2026 年 8 月<")
                data = text.encode("utf-8")
            target.writestr(item, data)
    temp.replace(path)


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(REFERENCE, OUTPUT)
    doc = Document(OUTPUT)
    remove_content_after_cover(doc)
    fill_cover(doc)
    add_content(doc)
    normalize_fonts(doc)
    props = doc.core_properties
    props.title = "调研模式前端框架重构计划书"
    props.subject = "在不改变调研链路的前提下重构独立双栏调研工作区"
    props.author = "Codex"
    props.comments = "基于现有代码框架与用户提供的 UI 参考图形成。"
    doc.save(OUTPUT)
    replace_header_placeholders(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
